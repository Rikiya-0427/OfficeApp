from docx import Document
from docx.shared import Inches

from docx import Document

# document = Document()

# document.add_heading('簡単なWordドキュメントのタイトル', 0)
# document.add_paragraph('簡単なWordドキュメントのテキスト')

# document.save('sample.docx')

# 「sample.docx」を読み込む
document = Document('sample.docx')

# 画像を貼り付ける
image_path = './images/keyboard-crasher-german-kid.gif'  # 貼り付ける画像のパス
document.add_picture(image_path, width=Inches(4), height=Inches(3))  # 画像のサイズを指定

# 文字数をカウントして出力する
text_count = 0
for paragraph in document.paragraphs:
    text_count += len(paragraph.text)

print("文字数:", text_count)

# 「sample_answer.docx」としてドキュメントを保存する
document.save('sample_answer.docx')